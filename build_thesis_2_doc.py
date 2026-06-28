import os
import subprocess
import sys

def convert_md_to_docx():
    print("Converting md to docx via markdown-docx...")
    subprocess.run('npx -y markdown-docx -i thesis_2_complete.md -o thesis_2.docx', shell=True, check=True)

def apply_styles():
    print("Applying styling formatting to docx...")
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx'], check=True)
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

    doc = Document('thesis_2.docx')
    
    # 1. Update document styles default font
    for style in doc.styles:
        if hasattr(style, 'font') and style.font is not None:
            style.font.name = 'Times New Roman'
            style.font.color.rgb = RGBColor(0, 0, 0)
            
    # 2. Iterate paragraphs to format text, sizes and alignment
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name.lower()
        
        # Paragraphs are justified (evenly distributed)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        for run in paragraph.runs:
            # Set run font name
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Explicitly set XML font attributes for compatibility
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            
            # Font size: 11 for normal paragraph/body text, larger for headings
            if 'heading 1' in style_name:
                run.font.size = Pt(18)
                run.font.bold = True
            elif 'heading 2' in style_name:
                run.font.size = Pt(15)
                run.font.bold = True
            elif 'heading 3' in style_name:
                run.font.size = Pt(13)
                run.font.bold = True
            elif 'heading 4' in style_name:
                run.font.size = Pt(12)
                run.font.bold = True
            else:
                run.font.size = Pt(11)

    # 3. Format tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        
                        rPr = run._r.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:ascii'), 'Times New Roman')
                        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    
    # Save both docx and doc versions
    doc.save('thesis_2.docx')
    doc.save('thesis_2.doc')

if __name__ == "__main__":
    convert_md_to_docx()
    apply_styles()
    print("All tasks completed successfully!")
