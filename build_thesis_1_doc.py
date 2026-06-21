import os
import re
import subprocess
import sys

def transform_chapter(content, chapter_num, section_index):
    """
    Transforms headings of a technical chapter to sit under 'III/ MATERIALS AND METHODS'.
    Example:
      # Chapter 1: AI Orchestration -> ## 3.1 AI Orchestration
      ## 1.1 Overview -> ### 3.1.1 Overview
      ### 1.2.1 Synchronous Inference -> #### 3.1.2.1 Synchronous Inference
    Also adjusts cross-reference mentions like §1.4.1 or §7.6.
    """
    # 1. Transform main chapter heading
    # Regex matches: # Chapter 1: Title or # Chapter 8: ML Training Pipeline — Title
    def main_heading_repl(match):
        title = match.group(1).strip()
        # Clean up any potential 'Chapter X:' or '—' artifacts
        title = re.sub(r'^(?:Chapter\s*\d+\s*:\s*|—\s*)', '', title)
        return f"## 3.{section_index} {title}"

    content = re.sub(r'^#\s+(.*)$', main_heading_repl, content, flags=re.MULTILINE)

    # 2. Transform subheadings
    # Matches: ## X.Y Title
    content = re.sub(
        rf'^##\s+{chapter_num}\.(\d+)\s+(.*)$',
        rf'### 3.{section_index}.\1 \2',
        content,
        flags=re.MULTILINE
    )

    # 3. Transform sub-subheadings
    # Matches: ### X.Y.Z Title
    content = re.sub(
        rf'^###\s+{chapter_num}\.(\d+)\.(\d+)\s+(.*)$',
        rf'#### 3.{section_index}.\1.\2 \3',
        content,
        flags=re.MULTILINE
    )

    # 4. Transform section references in body text (e.g., §1.4.1 -> §3.1.4.1)
    def ref_repl(match):
        chap = int(match.group(1))
        # Find which section_index corresponds to chap
        chapter_to_section = {1: 1, 3: 2, 6: 3, 7: 4, 8: 5}
        sec_idx = chapter_to_section.get(chap, chap)
        
        # If it has sub and subsub
        sub = match.group(2) or ""
        subsub = match.group(3) or ""
        
        if subsub:
            return f"§3.{sec_idx}.{sub.strip('.')}.{subsub.strip('.')}"
        elif sub:
            return f"§3.{sec_idx}.{sub.strip('.')}"
        else:
            return f"§3.{sec_idx}"

    # Regex matches: §1.4.1 or §3.2 or §7.6 or §8.2.1
    content = re.sub(r'§(\d+)\.(\d+)(?:\.(\d+))?', ref_repl, content)

    return content

def prepare_markdown():
    print("Preparing and combining markdown files...")
    
    # 1. Read individual files
    with open('thesis_1_front_matter.md', 'r', encoding='utf-8') as f:
        front_matter = f.read()
    with open('thesis_1_abstract.md', 'r', encoding='utf-8') as f:
        abstract = f.read()
    with open('thesis_1_introduction.md', 'r', encoding='utf-8') as f:
        introduction = f.read()
    with open('thesis_1_objectives.md', 'r', encoding='utf-8') as f:
        objectives = f.read()
        
    # Technical chapters to be organized under III/ MATERIALS AND METHODS
    technical_chapters = [
        ('chapter_1_thesis_1.md', 1, 1),
        ('chapter_3_thesis_1.md', 3, 2),
        ('chapter_6_thesis_1.md', 6, 3),
        ('chapter_7_thesis_1.md', 7, 4),
        ('chapter_8_thesis_1.md', 8, 5)
    ]
    
    methods_sections = ["# III/ MATERIALS AND METHODS\n"]
    for file_path, chap_num, sec_idx in technical_chapters:
        if not os.path.exists(file_path):
            print(f"Error: Required file {file_path} is missing!")
            sys.exit(1)
        with open(file_path, 'r', encoding='utf-8') as f:
            raw_content = f.read()
            transformed = transform_chapter(raw_content, chap_num, sec_idx)
            methods_sections.append(transformed)
            
    combined_methods = '\n\n'.join(methods_sections)
    
    with open('thesis_1_results_discussion.md', 'r', encoding='utf-8') as f:
        results_discussion = f.read()
    with open('thesis_1_conclusion.md', 'r', encoding='utf-8') as f:
        conclusion = f.read()
    with open('thesis_1_references.md', 'r', encoding='utf-8') as f:
        references = f.read()
    with open('thesis_1_appendices.md', 'r', encoding='utf-8') as f:
        appendices = f.read()
        
    # Assemble the final document parts in order
    thesis_parts = [
        front_matter,
        abstract,
        introduction,
        objectives,
        combined_methods,
        results_discussion,
        conclusion,
        references,
        appendices
    ]
    
    # Save combined file for conversion
    with open('thesis_1_complete.md', 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(thesis_parts))
    
    # Also save as thesis_1_with_images.md for backward compatibility
    with open('thesis_1_with_images.md', 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(thesis_parts))
        
    print("Successfully assembled markdown files into thesis_1_complete.md.")

def convert_md_to_docx():
    print("Converting md to docx via markdown-docx...")
    subprocess.run('npx -y markdown-docx -i thesis_1_complete.md -o thesis_1.docx', shell=True, check=True)

def apply_styles():
    print("Applying styling formatting to docx...")
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx'], check=True)
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

    doc = Document('thesis_1.docx')
    
    # 1. Update document styles default font
    for style in doc.styles:
        if hasattr(style, 'font') and style.font is not None:
            style.font.name = 'Times New Roman'
            style.font.color.rgb = RGBColor(0, 0, 0)
            
    # 2. Iterate paragraphs to format text, sizes and alignment
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name.lower()
        
        # Paragraphs are justified (evenly distributed)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        for run in paragraph.runs:
            # Set run font name
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Explicitly set XML font attributes for compatibility
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            
            # Font size: 11 for normal paragraph/body text, larger for headings
            if 'heading 1' in style_name:
                run.font.size = Pt(18)
                run.font.bold = True
            elif 'heading 2' in style_name:
                run.font.size = Pt(15)
                run.font.bold = True
            elif 'heading 3' in style_name:
                run.font.size = Pt(13)
                run.font.bold = True
            elif 'heading 4' in style_name:
                run.font.size = Pt(12)
                run.font.bold = True
            else:
                run.font.size = Pt(11)

    # 3. Format tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        
                        rPr = run._r.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:ascii'), 'Times New Roman')
                        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    
    # Save both docx and doc versions
    doc.save('thesis_1.docx')
    doc.save('thesis_1.doc')

if __name__ == "__main__":
    prepare_markdown()
    convert_md_to_docx()
    apply_styles()
    print("All tasks completed successfully!")
